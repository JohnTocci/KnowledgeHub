"""
File processing module for KnowledgeHub.
Handles various file types: PDF, DOCX, images, XLSX/CSV, etc.
"""
import os
import io
import re
import logging
from typing import Dict, List, Optional, Tuple, Any
from datetime import datetime
import PyPDF2
from docx import Document
import openpyxl
import pandas as pd
from PIL import Image
import base64

class FileProcessor:
    """Processes different file types and extracts content for summarization."""
    
    def __init__(self):
        self.supported_types = {
            'pdf': ['.pdf'],
            'document': ['.docx', '.doc'],
            'spreadsheet': ['.xlsx', '.xls', '.csv'],
            'image': ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp'],
            'text': ['.txt', '.md', '.markdown']
        }
    
    def is_supported_file(self, file_path: str) -> bool:
        """Check if file type is supported."""
        ext = os.path.splitext(file_path)[1].lower()
        return any(ext in exts for exts in self.supported_types.values())
    
    def get_file_type(self, file_path: str) -> str:
        """Determine file type category."""
        ext = os.path.splitext(file_path)[1].lower()
        
        for file_type, extensions in self.supported_types.items():
            if ext in extensions:
                return file_type
        
        return 'unknown'
    
    def process_file(self, file_path: str, title: str = None) -> Dict[str, Any]:
        """Process uploaded file and extract content."""
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")
            
            file_type = self.get_file_type(file_path)
            file_size = os.path.getsize(file_path)
            
            # Generate title if not provided
            if not title:
                title = os.path.splitext(os.path.basename(file_path))[0]
                title = self._clean_title(title)
            
            result = {
                'title': title,
                'file_path': file_path,
                'file_type': file_type,
                'file_size': file_size,
                'processed_date': datetime.now(),
                'content': '',
                'metadata': {},
                'images': [],
                'error': None
            }
            
            # Process based on file type
            if file_type == 'pdf':
                result.update(self._process_pdf(file_path))
            elif file_type == 'document':
                result.update(self._process_document(file_path))
            elif file_type == 'spreadsheet':
                result.update(self._process_spreadsheet(file_path))
            elif file_type == 'image':
                result.update(self._process_image(file_path))
            elif file_type == 'text':
                result.update(self._process_text(file_path))
            else:
                result['error'] = f"Unsupported file type: {file_type}"
            
            return result
            
        except Exception as e:
            logging.error(f"Error processing file {file_path}: {e}")
            return {
                'title': title or 'Unknown',
                'file_path': file_path,
                'file_type': 'unknown',
                'file_size': 0,
                'processed_date': datetime.now(),
                'content': '',
                'metadata': {},
                'images': [],
                'error': str(e)
            }
    
    def _process_pdf(self, file_path: str) -> Dict[str, Any]:
        """Process PDF file and extract text content."""
        try:
            content = ""
            metadata = {}
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract metadata
                if pdf_reader.metadata:
                    metadata.update({
                        'author': pdf_reader.metadata.get('/Author', ''),
                        'creator': pdf_reader.metadata.get('/Creator', ''),
                        'producer': pdf_reader.metadata.get('/Producer', ''),
                        'subject': pdf_reader.metadata.get('/Subject', ''),
                        'title': pdf_reader.metadata.get('/Title', ''),
                        'creation_date': pdf_reader.metadata.get('/CreationDate', ''),
                        'modification_date': pdf_reader.metadata.get('/ModDate', '')
                    })
                
                metadata['page_count'] = len(pdf_reader.pages)
                
                # Extract text from all pages
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            content += f"\n\n--- Page {page_num + 1} ---\n\n"
                            content += page_text
                    except Exception as e:
                        logging.warning(f"Error extracting text from page {page_num + 1}: {e}")
                        continue
            
            # Clean up content
            content = self._clean_extracted_text(content)
            
            return {
                'content': content,
                'metadata': metadata,
                'word_count': len(content.split()) if content else 0
            }
            
        except Exception as e:
            logging.error(f"Error processing PDF: {e}")
            return {
                'content': '',
                'metadata': {},
                'word_count': 0,
                'error': f"PDF processing error: {str(e)}"
            }
    
    def _process_document(self, file_path: str) -> Dict[str, Any]:
        """Process DOCX/DOC file and extract text content."""
        try:
            content = ""
            metadata = {}
            
            if file_path.lower().endswith('.docx'):
                doc = Document(file_path)
                
                # Extract metadata
                props = doc.core_properties
                metadata.update({
                    'author': props.author or '',
                    'title': props.title or '',
                    'subject': props.subject or '',
                    'keywords': props.keywords or '',
                    'created': str(props.created) if props.created else '',
                    'modified': str(props.modified) if props.modified else '',
                    'last_modified_by': props.last_modified_by or ''
                })
                
                # Extract text from paragraphs
                paragraphs = []
                for paragraph in doc.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        paragraphs.append(text)
                
                content = '\n\n'.join(paragraphs)
                
                # Extract text from tables
                for table in doc.tables:
                    table_text = self._extract_table_text(table)
                    if table_text:
                        content += f"\n\n--- Table ---\n\n{table_text}"
                
            else:
                # For .doc files, we'd need python-docx2txt or similar
                # For now, return an error message
                return {
                    'content': '',
                    'metadata': {},
                    'word_count': 0,
                    'error': '.doc files not supported, please convert to .docx'
                }
            
            # Clean up content
            content = self._clean_extracted_text(content)
            
            return {
                'content': content,
                'metadata': metadata,
                'word_count': len(content.split()) if content else 0
            }
            
        except Exception as e:
            logging.error(f"Error processing document: {e}")
            return {
                'content': '',
                'metadata': {},
                'word_count': 0,
                'error': f"Document processing error: {str(e)}"
            }
    
    def _process_spreadsheet(self, file_path: str) -> Dict[str, Any]:
        """Process Excel/CSV file and extract data summary."""
        try:
            content = ""
            metadata = {}
            
            if file_path.lower().endswith('.csv'):
                # Process CSV file
                df = pd.read_csv(file_path)
                
                metadata.update({
                    'rows': len(df),
                    'columns': len(df.columns),
                    'file_size_mb': round(os.path.getsize(file_path) / 1024 / 1024, 2)
                })
                
                # Generate summary
                content = self._generate_dataframe_summary(df)
                
            else:
                # Process Excel file
                workbook = openpyxl.load_workbook(file_path, read_only=True)
                sheet_names = workbook.sheetnames
                
                metadata.update({
                    'sheet_count': len(sheet_names),
                    'sheet_names': sheet_names,
                    'file_size_mb': round(os.path.getsize(file_path) / 1024 / 1024, 2)
                })
                
                # Process each sheet
                for sheet_name in sheet_names[:3]:  # Limit to first 3 sheets
                    try:
                        df = pd.read_excel(file_path, sheet_name=sheet_name)
                        sheet_summary = self._generate_dataframe_summary(df, sheet_name)
                        content += f"\n\n--- Sheet: {sheet_name} ---\n\n{sheet_summary}"
                    except Exception as e:
                        logging.warning(f"Error processing sheet {sheet_name}: {e}")
                        continue
                
                workbook.close()
            
            return {
                'content': content,
                'metadata': metadata,
                'word_count': len(content.split()) if content else 0
            }
            
        except Exception as e:
            logging.error(f"Error processing spreadsheet: {e}")
            return {
                'content': '',
                'metadata': {},
                'word_count': 0,
                'error': f"Spreadsheet processing error: {str(e)}"
            }
    
    def _process_image(self, file_path: str) -> Dict[str, Any]:
        """Process image file and extract metadata."""
        try:
            content = ""
            metadata = {}
            
            with Image.open(file_path) as img:
                metadata.update({
                    'format': img.format,
                    'mode': img.mode,
                    'size': img.size,
                    'width': img.width,
                    'height': img.height,
                    'has_transparency': 'transparency' in img.info or 'A' in img.mode
                })
                
                # Extract EXIF data if available
                if hasattr(img, '_getexif') and img._getexif():
                    exif_data = img._getexif()
                    if exif_data:
                        metadata['exif'] = dict(exif_data)
                
                # Generate description for AI processing
                content = f"""Image Analysis:
- Format: {img.format}
- Dimensions: {img.width} x {img.height} pixels
- Color Mode: {img.mode}
- File Size: {round(os.path.getsize(file_path) / 1024, 2)} KB

This image has been uploaded for analysis and categorization. AI processing can be used to:
- Identify objects, scenes, or text in the image
- Generate descriptive tags
- Categorize the image type
- Extract any visible text (OCR)
"""
            
            return {
                'content': content,
                'metadata': metadata,
                'word_count': len(content.split()) if content else 0,
                'images': [file_path]  # Include the image itself
            }
            
        except Exception as e:
            logging.error(f"Error processing image: {e}")
            return {
                'content': '',
                'metadata': {},
                'word_count': 0,
                'error': f"Image processing error: {str(e)}"
            }
    
    def _process_text(self, file_path: str) -> Dict[str, Any]:
        """Process text/markdown file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            metadata = {
                'encoding': 'utf-8',
                'lines': len(content.split('\n')),
                'characters': len(content)
            }
            
            return {
                'content': content,
                'metadata': metadata,
                'word_count': len(content.split()) if content else 0
            }
            
        except Exception as e:
            logging.error(f"Error processing text file: {e}")
            return {
                'content': '',
                'metadata': {},
                'word_count': 0,
                'error': f"Text processing error: {str(e)}"
            }
    
    def _generate_dataframe_summary(self, df: pd.DataFrame, sheet_name: str = None) -> str:
        """Generate a comprehensive summary of a pandas DataFrame."""
        try:
            summary_parts = []
            
            if sheet_name:
                summary_parts.append(f"Sheet: {sheet_name}")
            
            # Basic info
            summary_parts.append(f"Dimensions: {df.shape[0]} rows × {df.shape[1]} columns")
            summary_parts.append(f"Columns: {', '.join(df.columns.astype(str).tolist())}")
            
            # Data types
            dtype_summary = df.dtypes.value_counts()
            summary_parts.append(f"Data Types: {dict(dtype_summary)}")
            
            # Missing values
            missing_values = df.isnull().sum()
            if missing_values.sum() > 0:
                missing_info = missing_values[missing_values > 0]
                summary_parts.append(f"Missing Values: {dict(missing_info)}")
            
            # Sample data (first few rows)
            if len(df) > 0:
                summary_parts.append("\nSample Data (first 5 rows):")
                sample_df = df.head(5).to_string(index=False, max_cols=8, max_colwidth=30)
                summary_parts.append(sample_df)
            
            # Statistical summary for numeric columns
            numeric_cols = df.select_dtypes(include=['number']).columns
            if len(numeric_cols) > 0:
                summary_parts.append("\nNumeric Column Statistics:")
                stats_df = df[numeric_cols].describe()
                summary_parts.append(stats_df.to_string())
            
            # Unique values for categorical columns
            categorical_cols = df.select_dtypes(include=['object', 'category']).columns
            if len(categorical_cols) > 0 and len(categorical_cols) <= 5:
                summary_parts.append("\nCategorical Column Unique Values:")
                for col in categorical_cols[:3]:  # Limit to first 3
                    unique_vals = df[col].nunique()
                    if unique_vals <= 20:
                        unique_list = df[col].unique()[:10]  # Show first 10
                        summary_parts.append(f"{col}: {list(unique_list)} (Total unique: {unique_vals})")
                    else:
                        summary_parts.append(f"{col}: {unique_vals} unique values")
            
            return '\n\n'.join(summary_parts)
            
        except Exception as e:
            logging.error(f"Error generating DataFrame summary: {e}")
            return f"Error generating summary: {str(e)}"
    
    def _extract_table_text(self, table) -> str:
        """Extract text from a docx table."""
        try:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)
                table_data.append(' | '.join(row_data))
            
            return '\n'.join(table_data)
            
        except Exception as e:
            logging.error(f"Error extracting table text: {e}")
            return ""
    
    def _clean_extracted_text(self, text: str) -> str:
        """Clean up extracted text by removing excessive whitespace and formatting."""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove multiple newlines
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        
        # Clean up special characters that might interfere with processing
        text = re.sub(r'[^\w\s\-\.\,\!\?\:\;\(\)\[\]\{\}\"\'\/\@\#\$\%\&\*\+\=\<\>\|\~\`\^\n]', '', text)
        
        return text.strip()
    
    def _clean_title(self, title: str) -> str:
        """Clean and format title."""
        # Remove file extensions
        title = re.sub(r'\.[^.]+$', '', title)
        
        # Replace underscores and hyphens with spaces
        title = re.sub(r'[_-]', ' ', title)
        
        # Clean up whitespace
        title = re.sub(r'\s+', ' ', title).strip()
        
        # Capitalize appropriately
        title = title.title()
        
        return title or "Untitled Document"
    
    def get_supported_extensions(self) -> List[str]:
        """Get list of all supported file extensions."""
        extensions = []
        for ext_list in self.supported_types.values():
            extensions.extend(ext_list)
        return sorted(extensions)